from flask import Blueprint,render_template,request,url_for,redirect,session,flash,get_flashed_messages,abort
from app.utils import login_required,employee_required
from app.models import Authentication,Employee,Comment,RoadMap
from app import db
from sqlalchemy.exc import IntegrityError
from collections import defaultdict
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


employee_bp = Blueprint("employee",__name__,url_prefix="/employee")

opportunity_stages = {"Lead":"orange","RFP":"yellow","Presentation":"yellow","Demo":"yellow","Contract":"yellowgreen","Won":"green","Lost":"red"}
modules = ["ERP","HCM","SCM","EPM","PaaS","Primaver","Payroll"]


@employee_bp.route("/")
@login_required
@employee_required
def home():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("home.html",state="employee",role="employee",auth_name=auth_name,auth=auth)

@employee_bp.route("/logout",methods=["POST","GET"])
@login_required
@employee_required
def logout():
    auth_name = Authentication.query.get(session["user_id"]).name
    if request.method == "POST":
        session.pop("role",None)
        session.pop("user_id",None)
        session.pop("email",None)
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("base.login"))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("logout.html",state="employee",role="employee",auth=auth)

@employee_bp.route("/select_road_map",methods=["POST","GET"])
@login_required
@employee_required
def select_road_map():
    auth = Authentication.query.get(session.get("user_id"))
    road_maps = RoadMap.query.filter(RoadMap.assigned_to.any(Authentication.id==auth.id))
    return render_template("select_road_map.html",auth=auth,role="employee",road_maps=road_maps,state="employee")


@employee_bp.route("/reports",methods=["POST","GET"])
@login_required
@employee_required
def reports():
    auth = Authentication.query.get(session.get("user_id"))
    road_maps = RoadMap.query.filter(RoadMap.assigned_to.any(Authentication.id == auth.id)).all()
    now = datetime.now()
    return render_template("reports.html",auth=auth,role="employee",state="employee",reports=reports,now=now,road_maps=road_maps)

@employee_bp.route("/road_maps", methods=["POST","GET"])
@login_required
@employee_required
def road_maps():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth =  Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    road_maps = RoadMap.query.filter(RoadMap.assigned_to.any(Authentication.id == auth.id)).all()
    return render_template("road_maps.html",auth=auth,role="employee",mode="See All",now=now,road_maps=road_maps,state="employee",modules=modules,opportunity_stages=opportunity_stages)

@employee_bp.route("/road_map_overview/<int:rm_id>", methods=["POST", "GET"])
@login_required
@employee_required
def road_map_overview(rm_id):
    road_map = RoadMap.query.get(rm_id)
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    if request.method == "POST":
        comment_text = request.form.get("comment_text", "").strip()
        if comment_text:
            new_comment = Comment(
                road_map_id=rm_id,
                author=auth.name,
                comment=comment_text,
                timestamp=datetime.now(),
            )
            db.session.add(new_comment)
            db.session.commit()

        return redirect(url_for("employee.road_map_overview", rm_id=rm_id))  # reload page

    # GET
    comments = Comment.query.filter_by(road_map_id=rm_id).order_by(Comment.timestamp).all()
    return render_template("road_map_overview.html",state='employee',auth=auth,role="employee",roadmap=road_map,now=now,comments=comments,opportunity_stages=opportunity_stages,modules=modules)

@employee_bp.route("/edit_road_map/<int:rm_id>", methods=["POST", "GET"])
@login_required
@employee_required
def edit_road_map(rm_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    participants = RoadMap.query.get(rm_id).assigned_to
    participant_ids = [user.id for user in participants]
    names_auth = Authentication.query.order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    road_map = RoadMap.query.get(rm_id)
    now = datetime.now()
    print(road_map.assigned_to)
    print(names_auth[0])
    if request.method == "POST":
        industry = request.form.get("industry")
        deal_size = request.form.get("deal_size")
        opportunity_stage = request.form.get("opportunity_stage")   
        customer_party = request.form.get("customer_party")
        road_map.industry = industry
        road_map.deal_size = deal_size
        road_map.opportunity_stage = opportunity_stage
        road_map.customer_party = customer_party
        db.session.commit()
        flash("Road Map Edited successfully","success")
        now = datetime.now()
        return redirect(url_for("employee.road_maps"))
    return render_template("edit_road_map.html",Title="EDIT OBJECTIVES",state='employee',auth=auth,role="employee",now=now,names_auth=names_auth,opportunity_stages=opportunity_stages,modules=modules,road_map=road_map)


@employee_bp.route("/download-report-word/<int:objective_id>")
def download_report_word(objective_id):
    doc = Document()
    doc.add_heading("Objective Report", 0)
    doc.add_paragraph(f"Objective ID: {objective_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Message"
    hdr[1].text = "Status"
    hdr[2].text = "Timestamp"

    # add data rows
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@employee_bp.route("/download-report/<int:objective_id>")
def download_report(objective_id):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # write content
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, "Objective Report")

    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, f"Objective ID: {objective_id}")
    p.drawString(50, height - 100, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p.save()
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.pdf", mimetype="application/pdf")
